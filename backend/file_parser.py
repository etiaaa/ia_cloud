import io
from PyPDF2 import PdfReader
from docx import Document
from openpyxl import load_workbook


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".xls", ".txt"}


def extract_text(filename: str, content: bytes) -> str:
    """Extrait le texte d'un fichier (PDF, Word, Excel, TXT)."""
    ext = _get_extension(filename)

    if ext == ".pdf":
        return _extract_pdf(content)
    elif ext == ".docx":
        return _extract_docx(content)
    elif ext in (".xlsx", ".xls"):
        return _extract_excel(content)
    elif ext == ".txt":
        return content.decode("utf-8", errors="replace")
    else:
        raise ValueError(f"Format non supporté : {ext}")


def is_supported(filename: str) -> bool:
    return _get_extension(filename) in SUPPORTED_EXTENSIONS


def _get_extension(filename: str) -> str:
    return "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""


def _extract_pdf(content: bytes) -> str:
    reader = PdfReader(io.BytesIO(content))
    texts = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            texts.append(text)
    return "\n".join(texts)


def _extract_docx(content: bytes) -> str:
    doc = Document(io.BytesIO(content))
    texts = []
    for para in doc.paragraphs:
        if para.text.strip():
            texts.append(para.text)
    # Aussi extraire les tableaux
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                texts.append(row_text)
    return "\n".join(texts)


def _extract_excel(content: bytes) -> str:
    wb = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    texts = []
    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        texts.append(f"[Feuille: {sheet_name}]")
        for row in sheet.iter_rows(values_only=True):
            row_values = [str(cell) for cell in row if cell is not None]
            if row_values:
                texts.append(" | ".join(row_values))
    wb.close()
    return "\n".join(texts)
